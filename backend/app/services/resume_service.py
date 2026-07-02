"""Resume handling: validation, text extraction (PDF/DOCX), and the
deterministic 'resume quality' signals that feed the Hiring Score."""
from __future__ import annotations

import io
import re

from app.core.config import settings
from app.core.errors import ValidationFailed
from app.domain.scoring import ResumeQualitySignals

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
ALLOWED_EXT = {".pdf", ".docx"}

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")
LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/\S+", re.IGNORECASE)
GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/\S+", re.IGNORECASE)
URL_RE = re.compile(r"https?://\S+|www\.\S+", re.IGNORECASE)
NAME_LINE_RE = re.compile(r"^[A-Z][A-Za-z.'-]*(?:\s+[A-Z][A-Za-z.'-]*){1,3}$")

ACTION_VERBS = {
    "built", "designed", "developed", "implemented", "created", "led", "managed",
    "analyzed", "optimized", "automated", "deployed", "engineered", "improved",
    "launched", "delivered", "architected", "reduced", "increased", "migrated",
    "integrated", "tested", "shipped", "collaborated", "researched",
}


def validate_upload(filename: str, content_type: str, size: int) -> None:
    name = (filename or "").lower()
    ext = name[name.rfind(".") :] if "." in name else ""
    if ext not in ALLOWED_EXT:
        raise ValidationFailed("Only PDF and DOCX resumes are supported.")
    if size <= 0:
        raise ValidationFailed("Uploaded file is empty.")
    if size > settings.max_resume_bytes:
        raise ValidationFailed(
            f"Resume exceeds the {settings.max_resume_mb} MB limit."
        )


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract resume text, appending any hyperlink URLs (e.g. LinkedIn/GitHub icon
    links) that don't otherwise appear as visible text so they're still detectable."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        text, links = _extract_pdf(file_bytes)
    elif name.endswith(".docx"):
        text, links = _extract_docx(file_bytes)
    else:
        raise ValidationFailed("Unsupported resume format.")

    extra_links = [l for l in links if l not in text]
    if extra_links:
        text += "\n\n[Detected hyperlinks]\n" + "\n".join(extra_links)
    return text.strip()


def _extract_pdf(file_bytes: bytes) -> tuple[str, list[str]]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        links: list[str] = []
        for page in reader.pages:
            for annot in page.annotations or []:
                try:
                    obj = annot.get_object()
                    uri = obj.get("/A", {}).get("/URI")
                    if uri:
                        links.append(str(uri))
                except Exception:
                    continue
        return text, links
    except Exception as exc:
        raise ValidationFailed(f"Could not read PDF: {exc}") from exc


def _extract_docx(file_bytes: bytes) -> tuple[str, list[str]]:
    from docx import Document

    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = [_paragraph_text_with_hyperlinks(p) for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        links = [
            rel.target_ref
            for rel in doc.part.rels.values()
            if "hyperlink" in rel.reltype
        ]
        return "\n".join(parts), links
    except Exception as exc:
        raise ValidationFailed(f"Could not read DOCX: {exc}") from exc


def _paragraph_text_with_hyperlinks(paragraph) -> str:
    """python-docx's `.text` skips runs nested inside <w:hyperlink> elements
    (e.g. a name or icon that's itself a clickable link) — walk the XML to
    include them."""
    texts = [n.text for n in paragraph._p.iter() if n.tag.endswith("}t") and n.text]
    return "".join(texts) if texts else paragraph.text


def extract_contact_hints(raw_text: str) -> dict:
    """Deterministic regex fallback for contact/link fields, used when the AI
    extraction misses something that's literally present in the text."""
    text = raw_text or ""
    email_match = EMAIL_RE.search(text)
    phone_match = PHONE_RE.search(text)
    linkedin_match = LINKEDIN_RE.search(text)
    github_match = GITHUB_RE.search(text)

    other_urls = [
        u for u in URL_RE.findall(text)
        if "linkedin.com" not in u.lower() and "github.com" not in u.lower()
    ]

    full_name = None
    for line in text.splitlines()[:8]:
        candidate = line.strip()
        if candidate and NAME_LINE_RE.match(candidate):
            full_name = candidate
            break

    return {
        "full_name": full_name,
        "email": email_match.group(0) if email_match else None,
        "phone": phone_match.group(0).strip() if phone_match else None,
        "linkedin": linkedin_match.group(0) if linkedin_match else None,
        "github": github_match.group(0) if github_match else None,
        "portfolio": other_urls[0] if other_urls else None,
    }


def compute_quality_signals(raw_text: str) -> ResumeQualitySignals:
    """Deterministic resume completeness heuristics (no AI)."""
    text = raw_text or ""
    low = text.lower()
    words = re.findall(r"[A-Za-z][A-Za-z'+-]*", text)
    verb_hits = sum(1 for w in words if w.lower() in ACTION_VERBS)

    def has_any(*keys: str) -> bool:
        return any(k in low for k in keys)

    return ResumeQualitySignals(
        has_skills=has_any("skill", "technical proficienc", "technolog"),
        has_projects=has_any("project"),
        has_experience=has_any("experience", "employment", "work history"),
        has_education=has_any("education", "b.tech", "bachelor", "degree", "university"),
        has_certifications=has_any("certification", "certificate", "certified"),
        has_contact=bool(EMAIL_RE.search(text) or PHONE_RE.search(text)),
        word_count=len(words),
        action_verb_count=verb_hits,
    )
